"""docx 渲染器：ArtifactSpec → 策略报告 Word 文档。

章节结构由 spec 决定；溯源以正文内 ev-* 标记 + 文末"溯源明细"附录呈现。
"""

from __future__ import annotations

import io
from collections.abc import Mapping
from datetime import UTC, datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from finance_agent.artifacts.spec import (
    ArtifactSpec,
    ChangepointTableBlock,
    HeadingBlock,
    NarrativeBlock,
    TableBlock,
)
from finance_agent.provenance import Evidence
from finance_agent.skills.loader import SkillInfo

SUPPORTED_BLOCKS = {"heading", "narrative", "table", "changepoint_table"}

_MUTED = RGBColor(0x6B, 0x77, 0x8C)
_FONT_LATIN = "Helvetica Neue"
_FONT_EA = "PingFang SC"
_DOCX_EA_LANG = "zh-CN"

_CP_KIND_LABELS = {
    "trend_up": "趋势拐头向上", "trend_down": "趋势拐头向下",
    "accel_up": "加速上涨", "accel_down": "加速下跌",
    "drawdown": "回撤确认", "rally": "反弹确认", "volume_spike": "量能异常",
}


class RenderError(RuntimeError):
    pass


def _get_or_add(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.insert(0, child)
    return child


def _set_rfonts(rfonts) -> None:
    for attr in ("w:asciiTheme", "w:eastAsiaTheme", "w:hAnsiTheme", "w:cstheme"):
        rfonts.attrib.pop(qn(attr), None)
    rfonts.set(qn("w:ascii"), _FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), _FONT_LATIN)
    rfonts.set(qn("w:cs"), _FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), _FONT_EA)
    rfonts.set(qn("w:hint"), "eastAsia")


def _set_style_fonts(doc: Document) -> None:
    styles = doc.styles.element
    doc_defaults = _get_or_add(styles, "w:docDefaults")
    rpr_default = _get_or_add(doc_defaults, "w:rPrDefault")
    rpr = _get_or_add(rpr_default, "w:rPr")
    _set_rfonts(_get_or_add(rpr, "w:rFonts"))

    for style_id in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        style = doc.styles[style_id]
        rpr = style.element.get_or_add_rPr()
        _set_rfonts(rpr.get_or_add_rFonts())

    theme_font_lang = doc.settings.element.find(qn("w:themeFontLang"))
    if theme_font_lang is None:
        theme_font_lang = OxmlElement("w:themeFontLang")
        doc.settings.element.append(theme_font_lang)
    theme_font_lang.set(qn("w:val"), "en-US")
    theme_font_lang.set(qn("w:eastAsia"), _DOCX_EA_LANG)


def _evidence_mark(paragraph, refs: list[str]) -> None:
    if not refs:
        return
    run = paragraph.add_run(f"［溯源：{'、'.join(refs)}］")
    run.font.size = Pt(8)
    run.font.color.rgb = _MUTED
    run.font.superscript = True


def _add_table(doc: Document, headers: list[str], rows: list[list[str]],
               caption: str = "", evidence_refs: list[str] | None = None) -> None:
    if caption:
        para = doc.add_paragraph()
        run = para.add_run(caption)
        run.bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for c, header in enumerate(headers):
        table.rows[0].cells[c].text = header
    for row in rows:
        cells = table.add_row().cells
        for c in range(len(headers)):
            cells[c].text = str(row[c]) if c < len(row) else ""
    if evidence_refs:
        _evidence_mark(doc.add_paragraph(), evidence_refs)


def render_docx(
    spec: ArtifactSpec,
    *,
    datasets: Mapping = {},
    evidence: Mapping[str, Evidence],
    skill: SkillInfo,
) -> bytes:
    if spec.kind != "docx":
        raise RenderError(f"kind 不匹配：期望 docx，得到 {spec.kind}")
    unsupported = {b.type for b in spec.blocks} - SUPPORTED_BLOCKS
    if unsupported:
        raise RenderError(f"docx 渲染器不支持 block 类型：{sorted(unsupported)}")

    doc = Document()
    _set_style_fonts(doc)
    doc.add_heading(spec.title, level=0)
    meta = doc.add_paragraph()
    run = meta.add_run(
        (spec.subtitle + "　" if spec.subtitle else "")
        + f"生成于 {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}"
        "　｜　本报告为研究复盘工具，不构成投资建议"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = _MUTED

    for block in spec.blocks:
        if isinstance(block, HeadingBlock):
            doc.add_heading(block.text, level=block.level - 1)
        elif isinstance(block, NarrativeBlock):
            for para_text in block.text.split("\n\n"):
                para = doc.add_paragraph(para_text.strip())
            _evidence_mark(para, block.evidence_refs)
        elif isinstance(block, TableBlock):
            _add_table(doc, block.headers, block.rows,
                       caption=block.caption, evidence_refs=block.evidence_refs)
        elif isinstance(block, ChangepointTableBlock):
            _add_table(
                doc,
                ["日期", "类型", "严重度", "触发规则", "数据窗口"],
                [
                    [cp.date, _CP_KIND_LABELS.get(cp.kind, cp.kind), f"{cp.severity}/3",
                     cp.rule, f"{cp.window[0]} → {cp.window[1]}"]
                    for cp in block.changepoints
                ],
                caption=block.caption,
            )

    doc.add_heading("附录：溯源明细", level=1)
    if evidence:
        for ev in evidence.values():
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"{ev.id}").bold = True
            para.add_run(
                f"［{ev.kind}］{ev.source_url}　抓取于 {ev.fetched_at}"
                + (f"　查询：{ev.query}" if ev.query else "")
            )
    else:
        doc.add_paragraph("（本报告未登记 evidence）")

    doc.core_properties.title = spec.title
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
